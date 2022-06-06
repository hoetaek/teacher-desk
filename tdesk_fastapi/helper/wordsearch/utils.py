import copy
import random
import re
import string
from enum import Enum
from random import shuffle

import hgtk
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from helper.wordsearch.difficulty_option import Difficulty, DifficultyOption
from helper.wordsearch.word_position import WordPosition


class Language(Enum):
    KOREAN = 1
    ENGLISH = 2


class PuzzleData:
    def __init__(
        self,
        words,
        puzzle_difficulty_option: DifficultyOption,
        is_uppercase: bool = False,
        is_hint_twist: bool = False,
    ):

        self.words = list(filter(None, words))

        self.lang: Language
        if hgtk.checker.is_hangul(words[0]):
            self.lang = Language.KOREAN
        else:
            self.lang = Language.ENGLISH
        if self.lang == Language.KOREAN and is_uppercase == True:
            raise ValueError("There is no uppercase in Korean")

        if is_uppercase:
            self.words = [i.upper() for i in self.words]
        
        self.is_uppercase = is_uppercase
        self.is_hint_twist = is_hint_twist

        self.set_difficulty_option(puzzle_difficulty_option)

    def set_difficulty_option(self, puzzle_difficulty_option: DifficultyOption):
        puzzle_difficulty_option.configure(self.words)
        self.__difficulty_option = puzzle_difficulty_option
        self.puzzle = [
            [0 for _ in range(puzzle_difficulty_option.width)]
            for _ in range(puzzle_difficulty_option.height)
        ]

    @property
    def width(self):
        return self.__difficulty_option.width

    @property
    def height(self):
        return self.__difficulty_option.height

    def make(self):
        self.__make_heading()
        while self.__make_puzzle() == False:
            self.__difficulty_option.resize_bigger()
            self.__empty_puzzle()
        self.__make_hint()

    def __make_heading(self):
        match self.lang:
            case Language.KOREAN:
                self.heading = "낱말 찾기"
            case Language.ENGLISH:
                self.heading = "Word Search"

    def __make_hint(self):
        self.hint = []
        for word in self.words:
            if self.is_hint_twist:
                match self.lang:
                    case Language.ENGLISH:
                        spelling = [i for i in word]
                        shuffle(spelling)
                        word = "".join(spelling)
                    case Language.KOREAN:
                        chosung_word = ""
                        for chr in word:
                            chosung = hgtk.letter.decompose(chr)[0]
                            chosung_word += chosung
                        word = chosung_word
            self.hint.append(word)

    def __make_puzzle(self):
        for word in sorted(self.words, key=len, reverse=True):
            entering_word_succeed = False
            try_num = 0
            max_try = 30
            while not entering_word_succeed:
                word_positions = WordPosition(
                    word,
                    self.__difficulty_option.width,
                    self.__difficulty_option.height,
                )
                (
                    x_direction,
                    y_direction,
                ) = self.__difficulty_option.get_direction_options()
                positions = word_positions.get_word_positions(
                    x_direction,
                    y_direction,
                )
                if self.__place_for_word_exists(word, positions):
                    self.__fill_word_in_puzzle(word, positions)
                    entering_word_succeed = True
                try_num += 1
                if try_num > max_try:
                    return False

        self.answer = copy.deepcopy(self.puzzle)
        self.__fill_random_letters()
        [print(i) for i in self.puzzle]
        return True

    def __fill_word_in_puzzle(self, word: str, word_positions: list[tuple[int, int]]):
        for letter, x, y in zip(word, *zip(*word_positions)):
            self.puzzle[y][x] = letter

    def __fill_random_letters(self):
        for i in range(self.__difficulty_option.height):
            for j in range(self.__difficulty_option.width):
                fill_alph = self.__get_random_letter()
                if self.puzzle[i][j] == 0:
                    self.puzzle[i][j] = fill_alph

    def __get_random_letter(self):
        letter_to_fill: str
        match self.lang:
            case Language.KOREAN:
                f = open("./helper/wordsearch/random_words.txt", "r")
                data = f.read()
                regex_f = r"[가-힣]+"
                search_target_f = data
                data = "".join(list(set(re.findall(regex_f, search_target_f))))
                source_letters = data
            case Language.ENGLISH:
                if not self.is_uppercase:
                    source_letters = string.ascii_lowercase
                else:
                    source_letters = string.ascii_uppercase

        source_letters = self.__difficulty_option.revise_source_letters(
            self.words, source_letters
        )
        letter_to_fill = random.choice(source_letters)
        return letter_to_fill

    def __place_for_word_exists(self, word, word_positions: list[tuple[int, int]]):
        for x, y, letter in zip(*zip(*word_positions), word):
            if not (0 == self.puzzle[y][x] or letter == self.puzzle[y][x]):
                return False
        return True

    def __empty_puzzle(self):
        self.puzzle = [
            [0 for _ in range(self.__difficulty_option.width)]
            for _ in range(self.__difficulty_option.height)
        ]


class Worksheet:
    def __init__(self, puzzle_data: PuzzleData):
        self.width = puzzle_data.width
        self.height = puzzle_data.height
        self.heading = puzzle_data.heading
        self.puzzle = puzzle_data.puzzle
        self.hint: list[str] = puzzle_data.hint
        self.answer = puzzle_data.answer
        self.__configure_settings()

    def __configure_settings(self):
        # write to docx file
        # Write to docx to puzzle.docx
        self.document = Document()
        # changing the page margins
        sections = self.document.sections
        for section in sections:
            # section.top_margin = Cm(1)
            # section.bottom_margin = Cm(0.8)
            section.left_margin = Cm(2.3)
            section.right_margin = Cm(2.3)

    def __add_heading(self, grade: str = "__", class_num: str = "__"):
        head = self.document.add_heading(self.heading, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        belong = f"{grade}학년 {class_num}반 이름: _______"
        para_belong = self.document.add_paragraph(belong)
        para_belong.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def __add_puzzle(self):
        puzzle_table = self.document.add_table(
            rows=self.height, cols=self.width, style="Table Grid"
        )
        puzzle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_height = 7200 / self.height
        for i, row in enumerate(puzzle_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(self.set_height))
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(5)
                cell.text = self.puzzle[i][j]
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def __add_hint_in_table(self):
        word_num = len(self.hint)
        if word_num <= 15:
            size = 5
        elif word_num <= 21:
            size = (word_num + 2) // 4
        else:
            size = 6
        hint_table = self.document.add_table(
            rows=(word_num + size - 1) // size, cols=size, style="Table Grid"
        )
        hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(hint_table.rows):
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "60")
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                index = i * size + j

                # 단어 수 만큼 반복하기
                if index < word_num:
                    for paragraph in cell.paragraphs:
                        run = paragraph.add_run(self.hint[index])
                        font = run.font
                        font.name = "Arial"
                        font.size = Pt(13)

                        #####가운데 정렬!!
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def write(self):
        self.__add_heading()
        self.__add_puzzle()
        self.document.add_paragraph()
        self.__add_hint_in_table()

    def write_answer(self):
        # 정답 파일 쓰기
        self.answ_doc = Document()
        answer_table = self.answ_doc.add_table(
            rows=self.height, cols=self.width, style="Table Grid"
        )
        answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(answer_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(self.set_height))
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(8)
                cell.text = str(self.answer[i][j])
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                    if cell.text == "0":
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tcPr.append(tcVAlign)

    def save(self, filename_or_stream):
        # todo get answer of puzzle too
        return self.document.save(filename_or_stream)

    @property
    def get_doc(self):
        return self.document


if __name__ == "__main__":
    english_words = [
        "hello",
        "python",
        "anything",
        "else",
        "witch",
        "campaign",
        "creed",
        "law",
        "small",
        "kidney",
        "basin",
        "theory",
        "refer",
        "cow",
        "twin",
        "quality",
        "wording",
        "punch",
        "perfume",
        "hemispher",
        "training",
        "triangle",
        "opera",
        "gravity",
        "feelings",
    ]
    korean_words = ["경찰관", "오늘의음식점", "낱말찾기퍼즐", "한국어", "민주주의", "헌법", "법원"]

    puzzle_difficulty_option = DifficultyOption(Difficulty.DIFFICULT)
    puzzle_data = PuzzleData(
        korean_words,
        puzzle_difficulty_option,
        is_uppercase=False,
        is_hint_twist=False,
    )
    puzzle_data.make()
    worksheet = Worksheet(puzzle_data)
    worksheet.write()
    worksheet.write_answer()
    worksheet.save("filename")
